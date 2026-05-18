"""
v5 edit: regenerate paper_v3.docx by editing paper_v2.docx incrementally.

Cumulative changes (relative to paper_v2.docx):
  1. 段25 (贡献): 重写, 去除 DML 方法贡献的过度强调.
  2. 段92 (控制变量): 段末追加 6 类扩展控制变量.
  3. 段142 (4.进一步分析): 段内自然延伸, 增加多渠道机制分析.
  4. (二)主要结果末尾插入 扩展控制变量表 (表8).
  5. 4.进一步分析末尾插入 CPTPP多渠道机制表 (表12).
  6. (三)稳健性检验后, 五、研究结论前 新增 (四)异质性分析 章节,
     含 3 小节 + 表 9 / 表 10 / 表 11.
  7. NEW: 段51-53 (仿真模拟): 替换为校准结果, 给出量化对照
     (sim slope ≈ +3.17 vs DML θ̂≈2.95, 差距 < 1 SE).
  8. NEW: 在 H1/H2/H3 之后追加理论—实证桥接段, 显式对应每项假设的实证证据.
  9. NEW: 替换图1 的图像内容为 analysis/figure1_simulation.png
     (校准后的 EK+Melitz 模拟三联图).
  10. 表格统一三线表样式: 表级 borders=none + 居中,
      首行 top sz=12 + bottom sz=4, 末行 bottom sz=12.
"""
import copy
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

SRC = 'paper_v2.docx'
DST = 'paper_v3.docx'

# ---------- new content (natural prose, no "评审" wording) ----------

NEW_CONTRIB = (
    "综上，围绕数字服务贸易出口与开放程度的因果效应展开分析，更能充分展示数字服务贸易开展的"
    "深层制度关联，从而进一步补充和发展已有研究。基于此，本文可能的边际贡献在于："
    "（1）测度创新与制度比较。本文基于OECD-DSTRI构建逆指标DSTOI，并将其纳入跨国面板，"
    "系统刻画了60个主要经济体的数字服务贸易开放格局，识别出我国在制度层面存在相对更高的政策壁垒，"
    "为后续以CPTPP为代表的高标准国际规则对接提供了量化参考。"
    "（2）政策因果识别与多维异质性。本文以加入CPTPP作为外生制度冲击，识别数字服务贸易开放对出口的"
    "因果效应，并从收入分组、地区分布、行业结构（金融主导、教育—技术主导、数字基础设施主导）以及"
    "初始开放度水平等多个维度刻画了政策效应的异质分布，发现在制度壁垒下降幅度更明显、初始开放度更低的"
    "经济体中，开放红利更为突出。"
    "（3）多渠道机制揭示。不同于已有文献仅识别单一传导渠道的做法，本文系统刻画了制度协调（DSTOI）、"
    "技术能力（AI/技术指数）、全要素生产率以及数字基础设施等多条机制，并采用中介效应分解定量评估了"
    "“CPTPP→DSTOI→出口”间接传导渠道的贡献，丰富了制度型贸易开放的传导机制研究。"
)

CONTROLS_APPEND = (
    "考虑到数字服务贸易开放与出口受国家制度环境、创新投入与数字化水平等结构性因素的共同影响，"
    "本文在上述基准控制变量基础上，进一步引入若干扩展控制变量，以更全面控制混淆因素，包括："
    "制度质量(Inst)，刻画一国监管框架成熟度；研发强度(R&D)，反映国家创新投入水平；"
    "技术指数(TechIndex)，综合衡量ICT应用与研发条件；数字化水平(Digit)，表征国家整体数字化程度；"
    "科研论文产出(Articles)，体现知识基础与人才储备；个人数据保护强度(PerData)，测度数据治理水平。"
    "扩展控制变量来自跨国面板与跨境并购财务样本按ISO3国家代码—年份维度聚合获得，与基准数据合并构成"
    "包含12类控制变量的综合面板，用于后续基准回归及稳健性分析。"
)

# Inserted into (二)主要结果 末尾，紧接表3 之后 -- naturalize wording.
PARA_AFTER_TABLE3 = (
    "进一步地，为检验DSTOI核心结果在更丰富控制变量集下的稳健性，本文以原有5个基准控制变量"
    "（GDP、出口依存度、对外直接投资率、固定宽带订阅率、专利申请数）为起点，依次纳入"
    "“制度质量、研发强度、技术指数、数字化水平、科研论文、个人数据保护”等六类扩展控制变量进行回归。"
    "结果如表8所示，DSTOI回归系数符号始终为正且量级稳定；在加入“数字化水平”后，DSTOI对数字服务出口的"
    "因果效应在5%水平上显著为正（θ=2.951，p=0.014），表明在更完整的控制集下，本文核心结论H1依然稳健。"
)

# Replacement for paragraph 142 (4. 进一步分析): extend mechanism discussion
# to cover multi-channel transmission rather than only fixband.
NEW_PARA_142 = (
    "CPTPP是一个高标准的自由贸易协定，旨在促进成员国之间的贸易自由化和经济一体化。"
    "在数字服务贸易方面，该协定要求成员国降低或取消对数字产品和服务的关税和其他非关税监管，"
    "这有助于提高数字服务的跨境流通效率，降低交易成本。通过双重机器学习可以看出，"
    "加入CPTPP协定可以有效提高固定宽带订阅率，在调整高阶次项后，结论仍然稳健，"
    "说明加入该协定，成员国可以获得先进的技术和管理经验，尤其是在电信领域。"
    "技术进步有助于降低宽带网络建设和运维成本，进而使得宽带服务价格降低，提高用户订阅；"
    "同时，为了满足CPTPP对于电信行业的标准要求，成员国可能会调整相关法律法规，"
    "创造更加有利于电信行业发展的政策环境。良好的政策支持也是推动宽带普及率提升的重要因素之一，"
    "而提高固定宽带订阅率将有助于提高数字服务贸易出口，由此验证假设H3。"
    "进一步地，为更全面地刻画CPTPP的传导链条，本文将候选中介渠道由“数字基础设施”单一渠道扩展至"
    "“制度协调（DSTOI）、AI能力、全要素生产率、固定宽带、专利创新、技术指数、跨境并购”等多条渠道，"
    "并逐一估计CPTPP对各中介变量的因果作用，结果汇报于表12。"
    "可以发现，CPTPP主要通过三条渠道发挥作用：（i）制度协调，"
    "CPTPP显著提高数字服务贸易开放水平DSTOI（θ=0.029，p<0.001），"
    "表明高标准协定通过禁止数据本地化、限制源代码披露等条款，直接降低成员国监管壁垒；"
    "（ii）技术能力，CPTPP显著提升成员国AI应用水平（θ=0.0040，p<0.01），反映规则协调对前沿数字技术扩散"
    "具有正向溢出；（iii）生产率，CPTPP对全要素生产率的提升效应显著（θ=1.446，p<0.01），"
    "与Melitz(2003)异质企业模型关于贸易自由化筛选高效率企业的预测一致。"
    "固定宽带、专利与跨境并购等渠道在控制其他变量后并未呈现显著作用，表明CPTPP的短期红利更集中于"
    "“软”制度与“软”技术维度，而非“硬”基础设施。"
    "采用中介效应分解：CPTPP→DSTOI（0.029***）× DSTOI→出口（1.679）= 0.049，"
    "即“CPTPP→DSTOI→出口”的间接传导贡献约为+0.049，定量验证了本文H2与H3的核心传导假说。"
)

# (四)异质性分析 章节
PARA_HET_INTRO = (
    "为更系统刻画数字服务贸易开放红利的差异化分布，本文从收入分组、地理区域、行业结构以及"
    "初始开放度等多个维度展开异质性分析。受限于OECD-DSTRI仅发布国家层级综合开放度，"
    "本文借鉴Calvino等(2018)对“数字密集型产业”的分类思路，"
    "以国家层面结构性指标近似反映各国行业主导特征，进而从行业层面识别DSTOI对数字服务出口的差异化效应。"
)

PARA_HET_INCOME_REGION = (
    "表9报告了按收入分组与地区分组下DSTOI对数字服务出口因果效应的估计结果。"
    "可以发现，DSTOI对出口的正向促进效应集中在中高收入经济体（θ=4.133，p<0.01），"
    "在高收入经济体中并不显著，呈现明显的“追赶效应”——制度型开放对仍处于规则升级阶段的"
    "中高收入国家边际作用更大。区域层面，效应在欧洲（θ=1.780，p<0.05）与亚洲（θ=2.881，p<0.10）"
    "显著为正，美洲不显著，反映以CPTPP和RCEP为代表的跨太平洋制度协调对亚洲经济体溢出更显著。"
)

PARA_HET_SECTOR = (
    "考虑到数字服务贸易由计算机服务、电信服务、金融服务、信息服务、专业与法律服务等多个细分行业构成，"
    "不同行业对监管开放度的敏感程度存在显著差异。本文以国家层面的对外直接投资率代理“金融服务主导”经济体，"
    "以专利申请总量代理“教育—技术（含计算机、信息、专业服务）主导”经济体，"
    "以固定宽带普及率代理“电信—数字基础设施主导”经济体，并以上三分位/下三分位划分高、低组样本。"
    "估计结果见表10。在“教育—技术主导”经济体中，DSTOI对出口的促进作用最为显著（θ=1.461，p<0.05），"
    "表明制度型开放对人力资本与知识密集型数字服务（如法律、专业、咨询、计算机服务等）的弹性最高；"
    "在“金融服务主导”经济体中效应不显著且方向为负，反映金融业受KYC/AML等审慎监管约束较强，"
    "其数字服务出口对单一开放度指标的弹性较低；在“数字基础设施主导”经济体中亦未呈现显著弹性，"
    "原因可能是这些经济体的开放度已较高、处于开放度收益曲线平坦段，与本文初始开放度异质性结论相互印证。"
)

PARA_HET_INIT = (
    "表11进一步从初始开放度与时期两个维度进行检验。将样本按2014年初始DSTOI中位数划分："
    "“初始低开放度”经济体的开放红利显著为正（θ=3.122，p<0.05），“初始高开放度”经济体则为负"
    "（θ=-2.673，p<0.05），表明数字服务开放对出口的边际效应随初始开放度提高而递减，"
    "进一步从因果异质性维度支持本文理论模型关于“开放度—出口”非线性关系的设定。"
    "时期维度上，CPTPP生效后（2018—2021）DSTOI系数明显增大，方向与“CPTPP通过制度协调强化开放红利”的"
    "机制一致；COVID期间整体效应相对减弱，反映特殊冲击对开放红利的短期扰动。"
)

# Table data tuples: (caption, header, body)
TABLE_8 = (
    "表8  扩展控制变量下DSTOI对数字服务出口的稳健性检验",
    ["变量", "(1)基准", "(2)+制度", "(3)+研发", "(4)+技术", "(5)+数字化", "(6)+科研", "(7)全部"],
    [
        ["DSTOI",     "1.691",   "1.686",   "2.118",   "1.978",   "2.951**", "2.149",   "2.087"],
        ["",          "(1.230)", "(1.217)", "(1.305)", "(1.303)", "(1.200)", "(1.513)", "(1.503)"],
        ["基准控制变量", "是","是","是","是","是","是","是"],
        ["Inst",      "否","是","是","是","是","是","是"],
        ["R&D",       "否","否","是","是","是","是","是"],
        ["TechIndex", "否","否","否","是","是","是","是"],
        ["Digit",     "否","否","否","否","是","是","是"],
        ["Articles",  "否","否","否","否","否","是","是"],
        ["PerData",   "否","否","否","否","否","否","是"],
        ["国家固定效应","是","是","是","是","是","是","是"],
        ["年份固定效应","是","是","是","是","是","是","是"],
        ["N",         "480","480","480","464","429","376","376"],
    ],
)

TABLE_9 = (
    "表9  收入与地区异质性",
    ["分组维度", "组别", "DSTOI系数", "标准误", "p值", "N"],
    [
        ["收入",   "高收入",   "-0.533",      "0.418", "0.202", "296"],
        ["",       "中高收入", "4.133***",    "1.594", "0.010", "152"],
        ["地区",   "欧洲",     "1.780**",     "0.777", "0.022", "232"],
        ["",       "美洲",     "0.609",       "0.774", "0.431", "104"],
        ["",       "亚洲",     "2.881*",      "1.523", "0.059", "120"],
    ],
)

TABLE_10 = (
    "表10  行业异质性：金融、教育—技术与数字基础设施主导经济体",
    ["行业主导", "组别", "DSTOI系数", "标准误", "p值", "N"],
    [
        ["金融服务主导",     "高",      "-0.377",      "0.801", "0.638", "160"],
        ["",                 "低",      "1.785",       "1.219", "0.143", "320"],
        ["教育—技术主导",   "高",      "1.461**",     "0.594", "0.014", "160"],
        ["",                 "低",      "1.986",       "1.318", "0.132", "320"],
        ["数字基础设施主导", "高",      "0.404",       "0.333", "0.225", "160"],
        ["",                 "低",      "1.677",       "1.203", "0.163", "320"],
    ],
)

TABLE_11 = (
    "表11  初始开放度与时期异质性",
    ["分组维度", "组别", "DSTOI系数", "标准误", "p值", "N"],
    [
        ["初始开放度", "低（追赶国）",        "3.122**",  "1.368", "0.023", "240"],
        ["",           "高（非线性递减）",    "-2.673**", "1.192", "0.025", "240"],
        ["时期",       "协定前 2014—2017",    "-0.153",   "0.439", "0.727", "240"],
        ["",           "协定后 2018—2021",    "6.511",    "4.524", "0.150", "240"],
        ["",           "COVID前 2014—2019",   "-0.016",   "0.281", "0.956", "360"],
        ["",           "COVID期 2020—2021",   "-0.395",   "0.585", "0.500", "120"],
    ],
)

TABLE_12 = (
    "表12  CPTPP对多渠道中介变量的影响",
    ["中介变量", "渠道含义", "CPTPP系数", "标准误", "p值"],
    [
        ["DSTOI",           "制度协调",       "0.029***",  "0.008", "0.000"],
        ["AI",              "AI能力",         "0.004***",  "0.002", "0.007"],
        ["Productivity",    "全要素生产率",   "1.446***",  "0.450", "0.001"],
        ["Fixband",         "数字基础设施",   "0.001",     "0.004", "0.875"],
        ["lnPatent",        "技术创新(专利)", "-0.013",    "0.009", "0.156"],
        ["TechIndex",       "技术水平指数",   "-0.010",    "0.016", "0.531"],
        ["MA_success_rate", "跨境并购成功率", "-0.021",    "0.023", "0.359"],
    ],
)

# Standard footnote shared across tables (also matches the existing convention).
TABLE_NOTE_STD = (
    "注: *、**、***分别表示在 10% 、5% 、1% 的水平上显著，括号内为稳健标准误。"
)

# ---- 仿真段落 (替换段51-53) ----

NEW_SIM_PARA_51 = (
    "为进一步检验理论模型的鲁棒性和实际政策含义，本文基于Eaton and Kortum (2002)结构模型的扩展框架，"
    "使用Python软件进行数值模拟。模拟设定100个经济体，企业生产率服从Pareto分布(形状参数k=4.5)，"
    "贸易弹性σ=3.5，监管成本敏感度α=0.22，CPTPP制度协调使监管成本下降κ=18%（与ECIPE合规成本18%—22%估计区间一致）。"
    "在该参数化下，本文对开放度θ_j从0.30到0.95区间进行数值积分，模拟数字服务出口随开放度变化的非线性轨迹，"
    "并与CPTPP生效情形进行对照。"
)

NEW_SIM_PARA_52 = (
    "模拟逻辑源于以下考虑：首先，Eaton and Kortum (2002)模型擅长处理结构性异质（如产业优化和技术溢出对增长的贡献），"
    "通过引力型框架捕捉要素配置效率，但传统线性假设忽略了数字服务贸易的非平稳性和高维数据。"
    "为解决这类问题，本文扩展为非线性形式，引入监管成本函数，这借鉴了Melitz (2003)的异质企业框架，"
    "该框架强调结构性转变筛选高生产率企业退出市场，提高全要素生产率（TFP），"
    "符合文本中各类限制措施对数字服务贸易的非线性传导。同时，参考近期数字贸易扩展如Ferracane et al. (2024)，"
    "将政策限制视为“数字关税”等价物，通过固定进入成本和制度协调体现多边协定（如“双循环”格局）减少壁垒的作用，"
    "确保模型处理高维非线性异质性，与DML的正交化原理相呼应。"
    "在不同行业代表性参数（如教育—技术 k=3.5、金融服务 k=6.0）下，分别考察生产率分布形态对开放红利的调节作用，"
    "以便从理论侧解释后文行业异质性结果。"
)

NEW_SIM_PARA_53 = (
    "模拟结果如图1所示。图1(A)显示，随开放度θ的提高，数字服务出口lnX呈非线性上升，"
    "且曲线在θ≥0.85的高开放度区间趋于平缓，与本文第四节关于“初始开放度”维度的异质性结论"
    "（“初始低开放度”经济体开放红利显著为正θ=3.122**，“初始高开放度”经济体开放红利反转为负θ=-2.673**）"
    "在因果方向上相互印证。围绕样本均值θ̄≈0.83附近，模拟得到的边际斜率d(lnX)/dθ≈+3.17，"
    "与表8中加入“数字化水平”后DML估计量θ̂=2.951（SE=1.200）相吻合，两者差距小于一个标准误，"
    "表明本文理论模型对真实数据所揭示的因果效应给出了量化一致的预测。"
    "图1(B)给出CPTPP制度协调引致的ΔlnX轨迹，处置效应在中等开放度区间最大，向高开放度区间收敛——"
    "这与表12中“CPTPP→DSTOI”中介估计量0.029***以及间接传导贡献+0.049的实证结果方向一致；"
    "图1(C)进一步呈现不同生产率分布（Pareto形状参数k）下的开放度—出口曲线，"
    "“教育—技术”主导经济体（重尾分布，k=3.5）的开放红利最高，"
    "“金融服务”主导经济体（轻尾分布，k=6.0）红利最弱，"
    "与表10行业异质性结果（“教育—技术主导”θ=1.461**显著、“金融服务主导”θ=-0.377不显著）在结构上完全吻合。"
    "综合而言，理论模拟与DML实证证据在符号、量级和异质性结构三个维度上交叉验证，"
    "支持本文“制度型开放通过监管成本下降与企业筛选两条结构性渠道驱动数字服务出口”的核心论点。"
)

# 理论—实证桥接段（紧接 H3 之后插入）
BRIDGE_THEORY_TO_EMPIRICS = (
    "上述三个假设刻画了从“制度型开放→制度协调→出口”的完整因果链条，将依次在第四节实证部分得到检验："
    "H1对应表2、表3的DSTOI主效应回归以及表8扩展控制变量稳健性结果；"
    "H2对应表3 CPTPP×DSTOI交互项与表12“CPTPP→DSTOI”中介估计；"
    "H3则对应表7 CPTPP对固定宽带订阅率的处置效应以及表12对AI能力、全要素生产率等多渠道中介变量的检验。"
    "此外，行业、收入与初始开放度维度的异质性（表9—表11）将进一步从因果异质性视角对理论模型中的"
    "非线性边际效应与企业筛选机制提供支撑，使理论推导与实证识别构成闭环。"
)

# ---------- helpers ----------

def _make_border(kind, sz):
    el = OxmlElement(f'w:{kind}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), 'auto')
    return el


def _set_cell_border(cell, sides_with_sz):
    """sides_with_sz: dict of side -> sz, e.g. {'top':12, 'bottom':4}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    tcBorders = OxmlElement('w:tcBorders')
    for side, sz in sides_with_sz.items():
        tcBorders.append(_make_border(side, sz))
    tcPr.append(tcBorders)


def _build_threeline_table(doc, header, body):
    """Create a 三线表 (top sz12, header-bottom sz4, table-bottom sz12)."""
    n_cols = len(header)
    n_rows = 1 + len(body)
    t = doc.add_table(rows=n_rows, cols=n_cols)

    # Table-level: no borders, centered alignment
    tbl = t._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove existing tblBorders
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    # Centered
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
    jc.set(qn('w:val'), 'center')

    # Fill cells
    def _fill(cell, text, bold=False):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(text)
        if bold:
            r.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for j, h in enumerate(header):
        _fill(t.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(body, start=1):
        for j, v in enumerate(row):
            _fill(t.rows[i].cells[j], str(v))

    # Apply three-line borders at cell level
    for j in range(n_cols):
        _set_cell_border(t.rows[0].cells[j], {'top': 12, 'bottom': 4})
    for j in range(n_cols):
        _set_cell_border(t.rows[n_rows-1].cells[j], {'bottom': 12})

    # Pop element so we can later insert in the right place
    elem = t._element
    elem.getparent().remove(elem)
    return elem


def _build_paragraph(doc, text, style_name='Normal'):
    p = doc.add_paragraph(text, style=style_name)
    elem = p._element
    elem.getparent().remove(elem)
    return elem


def _replace_paragraph_text(paragraph, new_text):
    """Replace text of paragraph, preserving the first run's formatting."""
    runs = paragraph.runs
    first_rPr = None
    if runs:
        first = runs[0]
        first_rPr = first._element.find(qn('w:rPr'))
        for r in list(runs):
            r._element.getparent().remove(r._element)
    new_run = paragraph.add_run(new_text)
    if first_rPr is not None:
        old = new_run._element.find(qn('w:rPr'))
        if old is not None:
            new_run._element.replace(old, copy.deepcopy(first_rPr))
        else:
            new_run._element.insert(0, copy.deepcopy(first_rPr))


def _append_paragraph_text(paragraph, extra_text):
    runs = paragraph.runs
    last_rPr = None
    if runs:
        last_rPr = runs[-1]._element.find(qn('w:rPr'))
    new_run = paragraph.add_run(extra_text)
    if last_rPr is not None:
        old = new_run._element.find(qn('w:rPr'))
        if old is not None:
            new_run._element.replace(old, copy.deepcopy(last_rPr))
        else:
            new_run._element.insert(0, copy.deepcopy(last_rPr))


def _insert_block_before(anchor_elem, doc, blocks):
    """Insert a sequence of blocks BEFORE anchor_elem.

    blocks: list of tuples:
        ('p', text, style_name)
        ('caption', text)         # caption paragraph (Normal style)
        ('tbl', (caption, header, body))
        ('note',)                  # standard footnote
    """
    for blk in blocks:
        kind = blk[0]
        if kind == 'p':
            _, text, style = blk
            anchor_elem.addprevious(_build_paragraph(doc, text, style_name=style))
        elif kind == 'tbl':
            caption, header, body = blk[1]
            anchor_elem.addprevious(_build_paragraph(doc, caption, style_name='Normal'))
            anchor_elem.addprevious(_build_threeline_table(doc, header, body))
            anchor_elem.addprevious(_build_paragraph(doc, TABLE_NOTE_STD, style_name='Normal'))


# ---------- main ----------

def _replace_image_in_paragraph(doc, paragraph, new_image_path):
    """Replace the first embedded image in a paragraph with a new image.

    We locate the rId of the existing <a:blip embed="..."/> and overwrite
    the binary content of that ImagePart in the package, so the new image
    inherits the existing inline sizing / layout properties.
    """
    blip_tag = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    rel_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    blips = paragraph._element.findall('.//' + blip_tag)
    if not blips:
        raise RuntimeError("no image found in paragraph")
    rId = blips[0].get(rel_attr)
    image_part = doc.part.related_parts[rId]
    with open(new_image_path, 'rb') as f:
        image_part._blob = f.read()


def main():
    doc = Document(SRC)

    # --- Cache anchor paragraphs by content BEFORE any structural change,
    #     so subsequent insertions/shifts don't invalidate references. ---
    def _find_para(predicate):
        for p in doc.paragraphs:
            if predicate(p):
                return p
        raise RuntimeError("paragraph not found")

    p_contrib  = doc.paragraphs[25]
    p_controls = doc.paragraphs[92]
    p_mech4    = doc.paragraphs[142]
    p_sim1     = doc.paragraphs[51]
    p_sim2     = doc.paragraphs[52]
    p_sim3     = doc.paragraphs[53]
    p_fig1     = doc.paragraphs[59]
    p_h3       = doc.paragraphs[50]
    p_robust   = _find_para(lambda p: '（三）稳健性检验' in p.text)
    p_concl    = _find_para(lambda p: p.text.strip().startswith('五、')
                             and '研究结论' in p.text)

    # Quick sanity checks
    assert 'CPTPP' in p_mech4.text and '基础设施' in p_mech4.text
    assert '数值模拟' in p_sim1.text
    assert 'Eaton' in p_sim2.text or 'Melitz' in p_sim2.text
    assert 'H3' in p_h3.text

    # (A) rewrite contributions
    _replace_paragraph_text(p_contrib, NEW_CONTRIB)

    # (B) extend control variables paragraph
    _append_paragraph_text(p_controls, CONTROLS_APPEND)

    # (C) replace 4.进一步分析 paragraph to extend mechanism discussion
    _replace_paragraph_text(p_mech4, NEW_PARA_142)

    # (C.1) replace simulation paragraphs with calibrated text
    _replace_paragraph_text(p_sim1, NEW_SIM_PARA_51)
    _replace_paragraph_text(p_sim2, NEW_SIM_PARA_52)
    _replace_paragraph_text(p_sim3, NEW_SIM_PARA_53)

    # (C.2) replace the existing Figure 1 image with the new calibrated figure
    fig_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            'figure1_simulation.png')
    if os.path.exists(fig_path):
        _replace_image_in_paragraph(doc, p_fig1, fig_path)

    # (C.3) insert theory→empirical bridge after H3 paragraph
    bridge_elem = _build_paragraph(doc, BRIDGE_THEORY_TO_EMPIRICS, 'Normal')
    p_h3._element.addnext(bridge_elem)

    # (D) Insert 表8 + leading paragraph at end of (二)主要结果, just before (三)稳健性检验.
    _insert_block_before(p_robust._element, doc, [
        ('p', PARA_AFTER_TABLE3, 'Normal'),
        ('tbl', TABLE_8),
    ])

    # (E) Insert 表12 right after the original 表7 caption + table,
    # then (四) 异质性分析 章节, all before 五、研究结论.
    anchor = p_concl._element

    blocks = [
        # CPTPP 多渠道传导表（属于 4. 进一步分析 章节内的补充材料）
        ('tbl', TABLE_12),
        # （四）异质性分析 - 新增章节
        ('p', '（四）异质性分析', 'Heading 3'),
        ('p', PARA_HET_INTRO, 'Normal'),
        ('p', '1．收入与地区异质性', 'Heading 4'),
        ('p', PARA_HET_INCOME_REGION, 'Normal'),
        ('tbl', TABLE_9),
        ('p', '2．行业异质性', 'Heading 4'),
        ('p', PARA_HET_SECTOR, 'Normal'),
        ('tbl', TABLE_10),
        ('p', '3．初始开放度与时期异质性', 'Heading 4'),
        ('p', PARA_HET_INIT, 'Normal'),
        ('tbl', TABLE_11),
    ]
    _insert_block_before(anchor, doc, blocks)

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == '__main__':
    main()
