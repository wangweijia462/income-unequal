"""
Targeted v6 edit on the user's revised paper_v3_user.docx:

  1. Fix abnormally-wide line spacing on formula paragraphs 34-37
     (and 107, 128 table-caption paragraphs).
  2. Re-build tables 9 / 10 / 11 / 12 as 竖排 (variables as column
     headers, coefficient + SE-in-parens stacked vertically, no p-value
     column — significance shown by *** stars only).

Tables 8 and the existing 表 2/3 are already in this vertical format so
they are left untouched.

Per user request: keep edits minimal, leave everything else alone.
"""
import copy
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

SRC = 'paper_v3_user.docx'
DST = 'paper_v3.docx'

# ---- replacement table data: variables as column headers ----

# (caption, list of (var_name, coeff_with_stars, se, N)) + extra index labels
TABLE_9 = {
    'caption': '表9  收入与地区异质性',
    'cols': [
        # (column header, coeff, se, N)
        ('高收入',     '-0.533',      '0.418', '296'),
        ('中高收入',   '4.133***',    '1.594', '152'),
        ('欧洲',       '1.780**',     '0.777', '232'),
        ('美洲',       '0.609',       '0.774', '104'),
        ('亚洲',       '2.881*',      '1.523', '120'),
    ],
}

TABLE_10 = {
    'caption': '表10  行业异质性：金融、教育—技术与数字基础设施主导经济体',
    'cols': [
        ('金融主导(高)',       '-0.377',  '0.801', '160'),
        ('金融主导(低)',       '1.785',   '1.219', '320'),
        ('教育—技术(高)',     '1.461**', '0.594', '160'),
        ('教育—技术(低)',     '1.986',   '1.318', '320'),
        ('数字基础设施(高)',   '0.404',   '0.333', '160'),
        ('数字基础设施(低)',   '1.677',   '1.203', '320'),
    ],
}

TABLE_11 = {
    'caption': '表11  初始开放度与时期异质性',
    'cols': [
        ('初始低开放度',  '3.122**',  '1.368', '240'),
        ('初始高开放度',  '-2.673**', '1.192', '240'),
        ('协定前',        '-0.153',   '0.439', '240'),
        ('协定后',        '6.511',    '4.524', '240'),
        ('COVID 前',      '-0.016',   '0.281', '360'),
        ('COVID 期',      '-0.395',   '0.585', '120'),
    ],
}

TABLE_12 = {
    'caption': '表12  CPTPP对多渠道中介变量的影响',
    'cols': [
        ('DSTOI',         '0.029***',  '0.008', '480'),
        ('AI',            '0.004***',  '0.002', '480'),
        ('Productivity',  '1.446***',  '0.450', '155'),
        ('Fixband',       '0.001',     '0.004', '480'),
        ('lnPatent',      '-0.013',    '0.009', '480'),
        ('TechIndex',     '-0.010',    '0.016', '464'),
        ('MA_success',    '-0.021',    '0.023', '441'),
    ],
}


# ---- helpers ----

def _make_border(kind, sz):
    el = OxmlElement(f'w:{kind}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), 'auto')
    return el


def _set_cell_border(cell, sides_with_sz):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side, sz in sides_with_sz.items():
        tcBorders.append(_make_border(side, sz))
    tcPr.append(tcBorders)


def _build_threeline_table_vertical(doc, header_label, spec, coef_row_label='DSTOI'):
    """Build a vertical 三线表 of the form:

      [header_label]   col1     col2     col3 ...
      DSTOI            est1***  est2     est3 ...
                       (se1)    (se2)    (se3)
      N                n1       n2       n3
    """
    cols = spec['cols']
    n_cols = 1 + len(cols)                # +1 for left header
    n_rows = 4                            # header, DSTOI, SE, N
    t = doc.add_table(rows=n_rows, cols=n_cols)

    # table-level: no borders + centered
    tblPr = t._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        t._element.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); tblPr.append(jc)
    jc.set(qn('w:val'), 'center')

    def fill(cell, text, bold=False):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = 1
        r = p.add_run(text)
        if bold:
            r.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Row 0 -- column headers
    fill(t.rows[0].cells[0], '变量', bold=True)
    for j, (name, *_ ) in enumerate(cols, start=1):
        fill(t.rows[0].cells[j], name, bold=True)
    # Row 1 -- coefficients
    fill(t.rows[1].cells[0], coef_row_label)
    for j, (_, coef, _, _) in enumerate(cols, start=1):
        fill(t.rows[1].cells[j], coef)
    # Row 2 -- SE in parens
    fill(t.rows[2].cells[0], '')
    for j, (_, _, se, _) in enumerate(cols, start=1):
        fill(t.rows[2].cells[j], f'({se})')
    # Row 3 -- N
    fill(t.rows[3].cells[0], 'N')
    for j, (_, _, _, n) in enumerate(cols, start=1):
        fill(t.rows[3].cells[j], n)

    # three-line borders
    for j in range(n_cols):
        _set_cell_border(t.rows[0].cells[j], {'top': 12, 'bottom': 4})
        _set_cell_border(t.rows[n_rows-1].cells[j], {'bottom': 12})

    elem = t._element
    elem.getparent().remove(elem)
    return elem


def _fix_paragraph_spacing(paragraph):
    """Reset spacing to single-line and clear before/after spacing surplus."""
    pf = paragraph.paragraph_format
    # python-docx exposes .line_spacing as float multiplier when lineRule=auto.
    # Setting to 1.0 sets w:line="240" lineRule="auto".
    pf.line_spacing = 1.0
    # Also strip any w:spacing custom values by removing then re-creating
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            # Force lineRule="auto" + line="240"
            sp.set(qn('w:line'), '240')
            sp.set(qn('w:lineRule'), 'auto')


def _replace_table(doc, table_obj, new_elem):
    """Replace an existing table element with new_elem (preserve position)."""
    old_elem = table_obj._element
    parent = old_elem.getparent()
    parent.replace(old_elem, new_elem)


def main():
    doc = Document(SRC)

    # ---- (1) fix wide-spaced formula paragraphs ----
    for i in [34, 35, 36, 37]:
        p = doc.paragraphs[i]
        _fix_paragraph_spacing(p)
        # Sanity: print
        # print(f"  fixed P[{i}] ls -> {p.paragraph_format.line_spacing}")

    # ---- (2) rebuild tables 9 / 10 / 11 / 12 as vertical 三线表 ----
    # The user's docx has the same table sequence; tables[8..11] correspond
    # to 表9..表12 (since indices 0..7 are 表1..表8).
    # Build new elements first (each uses doc.add_table internally then is
    # detached), then replace in place.
    new_t9  = _build_threeline_table_vertical(doc, '变量', TABLE_9)
    new_t10 = _build_threeline_table_vertical(doc, '变量', TABLE_10)
    new_t11 = _build_threeline_table_vertical(doc, '变量', TABLE_11)
    new_t12 = _build_threeline_table_vertical(doc, '变量', TABLE_12,
                                              coef_row_label='CPTPP')

    _replace_table(doc, doc.tables[8],  new_t9)
    _replace_table(doc, doc.tables[9],  new_t10)
    _replace_table(doc, doc.tables[10], new_t11)
    _replace_table(doc, doc.tables[11], new_t12)

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == '__main__':
    main()
